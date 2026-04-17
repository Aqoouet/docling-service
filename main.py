"""Docling Service — converts DOCX/PDF to structured DoclingDocument JSON.

Single responsibility: receive a file upload, convert it with Docling,
return the DoclingDocument as JSON. All mapping to application-specific
models happens in the calling service (report_checking backend).

Note on Russian heading styles
-------------------------------
Docling's Word backend recognises only standard English style names
("Heading 1", "Heading 2", …). Documents created in Russian-locale Word
use localised names ("Заголовок 1", "Заголовок 2", …).  Before passing
the file to Docling we normalise any localised heading style names to
their English equivalents using python-docx.  The normalised file is
written to a second temp path and removed after conversion.
"""

from __future__ import annotations

import io
import os
import shutil
import tempfile
from pathlib import Path

from docling.datamodel.base_models import InputFormat
from docling.datamodel.pipeline_options import PdfPipelineOptions
from docling.document_converter import DocumentConverter, PdfFormatOption, WordFormatOption
from docling.pipeline.simple_pipeline import SimplePipeline
from docx import Document as DocxDocument
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.responses import JSONResponse

# ---------------------------------------------------------------------------
# Converter initialisation — done once at startup
# ---------------------------------------------------------------------------

_PDF_PIPELINE_OPTIONS = PdfPipelineOptions(do_ocr=False, do_table_structure=True)

_converter = DocumentConverter(
    format_options={
        # DOCX: lightweight rule-based pipeline, no ML models required
        InputFormat.DOCX: WordFormatOption(pipeline_cls=SimplePipeline),
        # PDF: full pipeline without OCR (text-based PDFs only)
        InputFormat.PDF: PdfFormatOption(pipeline_options=_PDF_PIPELINE_OPTIONS),
    }
)

# ---------------------------------------------------------------------------
# Localised → English heading style mapping
# ---------------------------------------------------------------------------

# Maps every known Russian/French/German/etc. heading style variant to the
# canonical English name that Docling's Word backend understands.
_LOCALISED_HEADING_MAP: dict[str, str] = {
    # Russian (Microsoft Word Russian locale)
    **{f"заголовок {i}": f"Heading {i}" for i in range(1, 10)},
    # German
    **{f"überschrift {i}": f"Heading {i}" for i in range(1, 10)},
    # French
    **{f"titre {i}": f"Heading {i}" for i in range(1, 10)},
}


def _normalise_docx_headings(src_path: str) -> str:
    """Return path to a copy of *src_path* with headings renamed to English.

    If no localised heading styles are found the original path is returned
    and no copy is made.
    """
    doc = DocxDocument(src_path)
    needs_fix = any(
        para.style.name.lower() in _LOCALISED_HEADING_MAP
        for para in doc.paragraphs
    )
    if not needs_fix:
        return src_path

    for para in doc.paragraphs:
        canonical = _LOCALISED_HEADING_MAP.get(para.style.name.lower())
        if canonical:
            try:
                para.style = doc.styles[canonical]
            except KeyError:
                # Style not present in this document — skip
                pass

    fixed_path = src_path + ".fixed.docx"
    doc.save(fixed_path)
    return fixed_path


# ---------------------------------------------------------------------------
# App
# ---------------------------------------------------------------------------

app = FastAPI(
    title="Docling Service",
    description="Converts DOCX/PDF documents to structured DoclingDocument JSON.",
    version="1.0.0",
)

_SUPPORTED_EXTENSIONS = {".docx", ".pdf"}


@app.get("/health")
def health() -> dict:
    return {"status": "ok"}


@app.post("/convert")
async def convert(file: UploadFile = File(...)) -> JSONResponse:
    """Convert an uploaded DOCX or PDF file to DoclingDocument JSON.

    Returns the full ``DoclingDocument.export_to_dict()`` payload which
    includes ``texts``, ``tables``, ``pictures``, and the document body tree.
    """
    filename = file.filename or "upload"
    suffix = Path(filename).suffix.lower()
    if suffix not in _SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{suffix}'. Supported: {sorted(_SUPPORTED_EXTENSIONS)}",
        )

    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    # Write to a temp file so Docling can open it by path
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    convert_path = tmp_path
    fixed_path: str | None = None
    try:
        if suffix == ".docx":
            convert_path = _normalise_docx_headings(tmp_path)
            if convert_path != tmp_path:
                fixed_path = convert_path

        result = _converter.convert(convert_path)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {exc}") from exc
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass
        if fixed_path:
            try:
                os.unlink(fixed_path)
            except OSError:
                pass

    doc_dict = result.document.export_to_dict()
    return JSONResponse(content=doc_dict)
